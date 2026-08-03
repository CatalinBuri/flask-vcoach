import os
import re
import json
import traceback
import httpx
from flask import Flask, request, jsonify, send_file
from io import BytesIO
from docx import Document

from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import mm
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer

app = Flask(__name__)
app.config["JSON_AS_ASCII"] = False


@app.after_request
def add_cors_headers(response):
    response.headers["Access-Control-Allow-Origin"] = "*"
    response.headers["Access-Control-Allow-Headers"] = "Content-Type,Authorization"
    response.headers["Access-Control-Allow-Methods"] = "GET,PUT,POST,DELETE,OPTIONS"
    return response


@app.before_request
def log_incoming_requests():
    print("\n--- [DIAGNOZA GLOBALA] Cerere primita ---", flush=True)
    print(f"Metoda: {request.method} | Path: {request.path}", flush=True)
    print(f"Antete (Headers): {dict(request.headers)}", flush=True)
    if request.method in ["POST", "PUT"]:
        if request.is_json:
            json_data = request.get_json(force=True, silent=True)
            print(f"Payload JSON primit: {json_data}", flush=True)
        elif request.form:
            print(f"Form data primit: {request.form.to_dict()}", flush=True)
        elif request.files:
            print(f"Fisiere primite: {list(request.files.keys())}", flush=True)
        else:
            print(f"Raw data / altele (lungime): {len(request.data)} bytes", flush=True)


MEMORY = {
    "cv_text": "",
    "job_description": "",
    "interview_history": [],
}


def call_mistral_api(
    prompt: str,
    model: str = "mistral-small-latest",
    temperature: float = 0.2,
    max_tokens: int = 4096,
) -> str:
    MISTRAL_API_KEY = os.environ.get("MISTRAL_API_KEY")
    if not MISTRAL_API_KEY:
        return ""

    try:
        url = "https://api.mistral.ai/v1/chat/completions"
        headers = {
            "Authorization": f"Bearer {MISTRAL_API_KEY}",
            "Content-Type": "application/json",
        }
        payload = {
            "model": model,
            "messages": [
                {
                    "role": "system",
                    "content": (
                        "Esti un asistent AI profesionist specializat in resurse umane, "
                        "optimizare CV-uri si interviuri."
                    ),
                },
                {"role": "user", "content": prompt},
            ],
            "temperature": temperature,
            "max_tokens": max_tokens,
        }

        with httpx.Client(timeout=30.0) as client:
            response = client.post(url, headers=headers, json=payload)
            response.raise_for_status()
            data = response.json()
            return data["choices"][0]["message"]["content"]
    except Exception as e:
        print(f"❌ Eroare API Mistral direct: {type(e).__name__} - {str(e)}", flush=True)
        return ""


gemini_client = None
GEMINI_API_KEY = os.environ.get("GEMINI_API_KEY")
MODEL_NAME = "gemini-2.5-flash"

if GEMINI_API_KEY:
    try:
        from google import genai

        gemini_client = genai.Client(api_key=GEMINI_API_KEY)
        print("✅ Gemini ready | model:", MODEL_NAME, flush=True)
    except Exception as e:
        print(f"⚠️ Gemini nu a putut fi initializat: {e}", flush=True)

groq_client = None
USE_GROQ = False
GROQ_API_KEY = os.environ.get("GROQ_API_KEY")

if GROQ_API_KEY:
    try:
        from groq import Groq

        groq_client = Groq(api_key=GROQ_API_KEY)
        USE_GROQ = True
        print("✅ Groq ready", flush=True)
    except Exception as e:
        print(f"⚠️ Groq nu a putut fi initializat: {e}", flush=True)

mistral_client = None
USE_MISTRAL = False
MISTRAL_API_KEY = os.environ.get("MISTRAL_API_KEY")

if MISTRAL_API_KEY:
    try:
        try:
            from mistralai import Mistral

            mistral_client = Mistral(api_key=MISTRAL_API_KEY)
            USE_MISTRAL = True
            print("✅ Mistral ready (SDK)", flush=True)
        except ImportError:
            try:
                from mistralai.client import MistralClient

                mistral_client = MistralClient(api_key=MISTRAL_API_KEY)
                USE_MISTRAL = True
                print("✅ Mistral ready (SDK Legacy)", flush=True)
            except ImportError:
                test_response = call_mistral_api("Spune 'test'")
                if "test" in test_response.lower():
                    USE_MISTRAL = True
                    print("✅ Mistral ready (Direct API)", flush=True)
                else:
                    print("⚠️ Mistral API key invalid sau conexiune esuata", flush=True)
    except Exception as e:
        print(f"⚠️ Mistral nu a putut fi initializat: {e}", flush=True)


def clean_text(text: str) -> str:
    if not text:
        return ""
    text = re.sub(r"\r\n|\r", "\n", text)
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def remove_consecutive_duplicates(text: str) -> str:
    if not text or not isinstance(text, str):
        return ""
    cleaned = re.sub(
        r"\b([a-zA-ZăâîșțĂÂÎȘȚ]+)(?:\s+\1\b)+",
        r"\1",
        text,
        flags=re.IGNORECASE,
    )
    cleaned = re.sub(r"(?i)\b([A-Zăâîșț\s]+)(\r?\n\1\b)+", r"\1", cleaned)
    return cleaned


def enforce_factuality_and_language(target_lang: str) -> str:
    if target_lang == "ro":
        lang_instruction = (
            "REGULA LINGVISTICA STRICTA: Tot outputul trebuie sa fie exclusiv in limba ROMANA. "
            "Nu amesteca limbi."
        )
    elif target_lang == "en":
        lang_instruction = (
            "STRICT LANGUAGE RULE: The entire output must be exclusively in ENGLISH. "
            "Do not mix languages."
        )
    else:
        lang_instruction = (
            "LANGUAGE RULE: Detect and use a single unified language consistently throughout."
        )

    anti_hallucination = (
        "REGULA ANTI-HALUCINATIE CRUCIALA: Este STRICT INTERZIS sa inventezi date, publicatii, "
        "companii sau experiente care nu exista in textul original furnizat de utilizator."
    )
    return f"{anti_hallucination}\n{lang_instruction}"


def safe_json(raw_text: str) -> dict:
    if not raw_text:
        return {}
    cleaned = re.sub(r"^```(?:json)?\s*", "", raw_text.strip(), flags=re.MULTILINE)
    cleaned = re.sub(r"\s*```$", "", cleaned, flags=re.MULTILINE).strip()

    try:
        return json.loads(cleaned)
    except Exception:
        match = re.search(r"\{.*\}", cleaned, re.DOTALL)
        if match:
            try:
                return json.loads(match.group(0))
            except Exception:
                pass
    return {}


def api_response(payload=None, error=None, code=200):
    if error:
        return (
            jsonify(
                {
                    "status": "error",
                    "success": False,
                    "ok": False,
                    "message": error,
                    "error": error,
                }
            ),
            code,
        )

    base_response = {
        "status": "success",
        "success": True,
        "ok": True,
        "code": 200,
    }

    if isinstance(payload, dict):
        base_response["data"] = payload
        base_response.update(payload)
        return jsonify(base_response), code

    base_response["data"] = payload if payload is not None else {}
    return jsonify(base_response), code


def gemini_text(prompt: str) -> str:
    if gemini_client:
        try:
            response = gemini_client.models.generate_content(
                model=MODEL_NAME,
                contents=prompt,
            )
            if response and hasattr(response, "text") and response.text:
                return response.text.strip()
        except Exception as e:
            print(f"⚠️ Eroare Gemini: {type(e)} - {str(e)}", flush=True)

    if USE_GROQ and groq_client:
        try:
            res = groq_client.with_options(max_retries=0).chat.completions.create(
                model="llama-3.1-8b-instant",
                messages=[
                    {"role": "system", "content": "Esti un asistent AI specializat in resurse umane."},
                    {"role": "user", "content": prompt},
                ],
                temperature=0.2,
                max_tokens=4096,
                timeout=12.0,
            )
            if res and res.choices and res.choices[0].message.content:
                return res.choices[0].message.content.strip()
        except Exception as e:
            print(f"⚠️ Eroare Groq: {type(e)} - {str(e)}", flush=True)

    if USE_MISTRAL:
        try:
            if mistral_client and hasattr(mistral_client, "chat"):
                res = mistral_client.chat.complete(
                    model="mistral-small-latest",
                    messages=[
                        {"role": "system", "content": "Esti un asistent AI specializat in resurse umane."},
                        {"role": "user", "content": prompt},
                    ],
                    temperature=0.2,
                    max_tokens=4096,
                )
                if res and res.choices and res.choices[0].message.content:
                    return res.choices[0].message.content.strip()

            direct_res = call_mistral_api(prompt)
            if direct_res:
                return direct_res
        except Exception as e:
            print(f"❌ Eroare Mistral: {type(e)} - {str(e)}", flush=True)
            return call_mistral_api(prompt)

    return ""


@app.route("/", methods=["GET", "HEAD", "OPTIONS"])
def index():
    if request.method == "OPTIONS":
        return api_response(code=200)
    return (
        jsonify(
            {
                "status": "online",
                "success": True,
                "service": "vCoach AI API",
                "gemini_active": gemini_client is not None,
                "groq_active": USE_GROQ,
                "mistral_active": USE_MISTRAL,
            }
        ),
        200,
    )


@app.route("/ping", methods=["GET", "HEAD", "OPTIONS"])
def ping():
    if request.method == "OPTIONS":
        return "", 200
    return "OK", 200


@app.route("/upload-cv", methods=["POST", "OPTIONS"], endpoint="upload_cv_root")
@app.route("/api/upload-cv", methods=["POST", "OPTIONS"], endpoint="upload_cv_api")
def upload_cv():
    if request.method == "OPTIONS":
        return api_response(code=200)

    try:
        text_content = ""
        if "file" in request.files:
            file = request.files["file"]
            filename = file.filename.lower()
            if filename.endswith(".pdf"):
                try:
                    import fitz

                    doc = fitz.open(stream=file.read(), filetype="pdf")
                    for page in doc:
                        text_content += page.get_text()
                except Exception as pdf_err:
                    return api_response(error=f"Eroare citire PDF: {str(pdf_err)}", code=400)
            else:
                text_content = file.read().decode("utf-8", errors="ignore")
        elif request.is_json:
            data = request.get_json(force=True, silent=True) or {}
            text_content = data.get("cv_text", "")

        cleaned = clean_text(text_content)
        if not cleaned:
            return api_response(error="Nu s-a putut extrage text din fisierul trimis.", code=400)

        MEMORY["cv_text"] = cleaned
        return api_response(
            payload={"message": "CV incarcat cu succes", "length": len(cleaned), "cv_text": cleaned}
        )
    except Exception as e:
        return api_response(error=f"Eroare la procesare: {str(e)}", code=500)


@app.route("/analyze-cv-quality", methods=["POST", "OPTIONS"], endpoint="analyze_cv_quality_root")
@app.route("/api/cv-quality", methods=["POST", "OPTIONS"], endpoint="analyze_cv_quality_api")
def analyze_cv_quality():
    if request.method == "OPTIONS":
        return api_response(code=200)

    try:
        data = request.get_json(force=True, silent=True) or {}
        cv_raw = data.get("cv_text") or MEMORY.get("cv_text") or ""
        job_raw = data.get("job_description") or data.get("job_text") or MEMORY.get("job_description") or ""
        target_lang = data.get("target_language") or data.get("language") or "ro"

        cv = clean_text(cv_raw)
        job = clean_text(job_raw)

        if not cv:
            return api_response(error="CV lipsa.", code=400)

        MEMORY["cv_text"] = cv
        if job:
            MEMORY["job_description"] = job

        factuality_rules = enforce_factuality_and_language(target_lang)

        if job:
            prompt = f"""
{factuality_rules}
Esti un recruiter senior si expert in sisteme ATS. Analizeaza CV-ul in raport direct cu Descrierea Jobului.
Raspunde EXCLUSIV cu un obiect JSON valid:
{{
  "clarity_score": 8,
  "relevance_score": 7,
  "structure_score": 8,
  "matched_ats_keywords": ["Cuvant1"],
  "missing_ats_keywords": ["CuvantLipseste"],
  "concrete_improvements": ["Sfat 1"],
  "suggested_rephrasings": ["Exemplu"]
}}
CV:
{cv}
DESCRIERE JOB:
{job}
"""
        else:
            prompt = f"""
{factuality_rules}
Esti un recruiter senior. Analizeaza structura si calitatea acestui CV.
Raspunde EXCLUSIV cu un obiect JSON valid:
{{
  "clarity_score": 8,
  "relevance_score": 6,
  "structure_score": 8,
  "detected_skills": ["Skill1"],
  "missing_ats_keywords": ["Adaugati un Job Description"],
  "concrete_improvements": ["Recomandare 1"],
  "suggested_rephrasings": ["Exemplu"]
}}
CV:
{cv}
"""

        raw_res = gemini_text(prompt)
        parsed = safe_json(raw_res)

        improvements = [
            remove_consecutive_duplicates(imp)
            for imp in (parsed.get("concrete_improvements") or [])
        ]
        rephrasings = [
            remove_consecutive_duplicates(rep)
            for rep in (parsed.get("suggested_rephrasings") or [])
        ]

        payload = {
            "clarity_score": parsed.get("clarity_score", 8),
            "relevance_score": parsed.get("relevance_score", 7 if job else 5),
            "structure_score": parsed.get("structure_score", 8),
            "has_job_context": bool(job),
            "ats_keywords": parsed.get("matched_ats_keywords")
            or parsed.get("detected_skills")
            or [],
            "missing_keywords": parsed.get("missing_ats_keywords") or [],
            "concrete_improvements": improvements,
            "suggested_rephrasings": rephrasings,
        }

        return api_response(payload=payload)
    except Exception as e:
        return api_response(error=f"Eroare analiza CV: {str(e)}", code=500)


@app.route("/interview-question", methods=["POST", "OPTIONS"], endpoint="interview_question_root")
@app.route("/api/interview-question", methods=["POST", "OPTIONS"], endpoint="interview_question_api")
def interview_question():
    if request.method == "OPTIONS":
        return api_response(code=200)

    try:
        data = request.get_json(force=True, silent=True) or {}
        user_answer = data.get("user_answer", "")
        role = data.get("role", "Software Developer")
        target_lang = data.get("target_language") or data.get("language") or "ro"

        factuality_rules = enforce_factuality_and_language(target_lang)
        prompt = f"""
{factuality_rules}
Esti un recrutator pentru rolul: {role}. Raspuns candidat: "{user_answer}"
Returneaza DOAR un obiect JSON valid:
{{
  "feedback": "Evaluare...",
  "score": 8,
  "next_question": "Urmatoarea intrebare..."
}}
"""
        raw_res = gemini_text(prompt)
        parsed = safe_json(raw_res) or {}

        feedback = remove_consecutive_duplicates(parsed.get("feedback", ""))
        next_q = remove_consecutive_duplicates(parsed.get("next_question", ""))

        payload = {
            "feedback": feedback,
            "score": parsed.get("score", 7),
            "next_question": next_q,
            "question": next_q,
        }

        return api_response(payload=payload)
    except Exception as e:
        return api_response(error=f"Eroare interviu: {str(e)}", code=500)


@app.route("/rephrase", methods=["POST", "OPTIONS"], endpoint="rephrase_root")
@app.route("/api/rephrase", methods=["POST", "OPTIONS"], endpoint="rephrase_api")
def rephrase():
    if request.method == "OPTIONS":
        return api_response(code=200)

    try:
        data = request.get_json(force=True, silent=True) or {}
        cv_text = data.get("text") or MEMORY.get("cv_text") or ""
        job_desc = data.get("job_description") or MEMORY.get("job_description") or ""
        target_lang = data.get("target_language") or data.get("language") or "ro"

        recommendations = data.get("recommendations") or data.get("concrete_improvements") or []
        missing_keywords = data.get("missing_keywords") or []
        matching_skills = data.get("matching_skills") or data.get("ats_keywords") or []

        if not cv_text:
            return api_response(error="Textul CV-ului pentru reformulare lipseste.", code=400)

        factuality_rules = enforce_factuality_and_language(target_lang)

        extra_context = ""
        if recommendations or missing_keywords:
            extra_context = f"""
SUGESTII DIN ANALIZA DE COMPATIBILITATE (integreaza-le natural, fara a inventa experiente):
- Recomandari concrete: {recommendations}
- Cuvinte cheie / skills lipsa: {missing_keywords}
- Skills deja potrivite: {matching_skills}
"""

        if job_desc:
            prompt = f"""
{factuality_rules}
Esti un expert in scriere de CV-uri si optimizare ATS.
Rescrie, structureaza si refocalizeaza complet continutul acestui CV bazandu-te exclusiv pe faptele reale din CV si aliniindu-l cu Descrierea Jobului.
Foloseste verbe puternice de actiune. Integreaza natural cuvintele cheie lipsa DOAR daca sunt sustinute de experienta reala din CV.
{extra_context}

Returneaza DOAR un obiect JSON valid cu structura:
{{
  "improved_text": "Textul complet rescris si optimizat al CV-ului..."
}}

CV ORIGINAL:
{cv_text}

DESCRIERE JOB:
{job_desc}
"""
        else:
            prompt = f"""
{factuality_rules}
Esti un expert in scriere de CV-uri. Imbunatateste si reformuleaza acest CV pe baza exclusiva a datelor reale existente.
{extra_context}

Returneaza DOAR un obiect JSON valid cu structura:
{{
  "improved_text": "Textul optimizat..."
}}

CV ORIGINAL:
{cv_text}
"""

        raw_res = gemini_text(prompt)
        parsed = safe_json(raw_res)
        improved = parsed.get("improved_text") if parsed else raw_res
        improved = remove_consecutive_duplicates(improved)

        payload = {
            "improved_text": improved,
            "rephrased_text": improved,
            "text": improved,
        }

        return api_response(payload=payload)
    except Exception as e:
        return api_response(error=f"Eroare rephrase: {str(e)}", code=500)


@app.route("/generate-cover-letter", methods=["POST", "OPTIONS"], endpoint="cover_letter_root")
@app.route("/api/cover-letter", methods=["POST", "OPTIONS"], endpoint="cover_letter_api")
def generate_cover_letter():
    if request.method == "OPTIONS":
        return api_response(code=200)

    try:
        # 1. Mai întâi citim datele din request
        data = request.get_json(force=True, silent=True) or {}
        
        # 2. Apoi extragem câmpurile
        company_name = (data.get("company_name") or "").strip()
        job_title = (data.get("job_title") or "").strip()

        if not company_name or not job_title:
            return api_response(
                error="Numele companiei și titlul jobului sunt obligatorii pentru Cover Letter.",
                code=400,
            )

        cv = clean_text(data.get("cv_text") or MEMORY.get("cv_text") or "")
        job_desc = clean_text(data.get("job_description") or MEMORY.get("job_description") or "")
        target_lang = data.get("target_language") or data.get("language") or "ro"

        if not cv or not job_desc:
            return api_response(
                error="CV-ul și Descrierea Jobului sunt necesare pentru Cover Letter.",
                code=400,
            )

        factuality_rules = enforce_factuality_and_language(target_lang)
        prompt = f"""
{factuality_rules}
Creează o scrisoare de intenție (Cover Letter) profesională, concisă (maximum 400 de cuvinte),
pentru rolul "{job_title}" la compania "{company_name}".

Folosește DOAR informații reale din CV. Nu inventa experiențe.

Structură:
1. Introducere – interes pentru rolul {job_title} la {company_name}
2. 1-2 paragrafe cu realizări relevante din CV
3. Încheiere – entuziasm, disponibilitate pentru interviu

CV:
{cv}

JOB DESCRIPTION:
{job_desc}
"""
        cover_letter_text = remove_consecutive_duplicates(gemini_text(prompt))
        payload = {
            "cover_letter": cover_letter_text,
            "text": cover_letter_text,
        }
        return api_response(payload=payload)
    except Exception as e:
        return api_response(error=f"Eroare cover letter: {str(e)}", code=500)


@app.route("/export-docx", methods=["POST", "OPTIONS"], endpoint="export_docx_root")
@app.route("/api/export-docx", methods=["POST", "OPTIONS"], endpoint="export_docx_api")
def export_docx():
    if request.method == "OPTIONS":
        return "", 200

    try:
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

        data = request.get_json(force=True, silent=True) or {}
        text_content = data.get("text") or MEMORY.get("cv_text") or ""

        print(f"--- [DIAGNOZA EXPORT DOCX] --- Lungime text: {len(text_content)} caractere", flush=True)

        if not text_content:
            return api_response(error="Text lipsa pentru export.", code=400)

        def add_runs_with_bold(paragraph, text):
            """Parseaza **bold** din text si adauga runs corespunzatoare."""
            import re
            parts = re.split(r'(\*\*.*?\*\*)', text)
            for part in parts:
                if part.startswith('**') and part.endswith('**') and len(part) > 4:
                    run = paragraph.add_run(part[2:-2])
                    run.bold = True
                else:
                    subparts = re.split(r'(_.*?_)', part)
                    for sp in subparts:
                        if sp.startswith('_') and sp.endswith('_') and len(sp) > 2:
                            run = paragraph.add_run(sp[1:-1])
                            run.italic = True
                        else:
                            paragraph.add_run(sp)

        doc = Document()

        style = doc.styles['Normal']
        style.font.name = 'Arial'
        style.font.size = Pt(10)

        for line in text_content.split("\n"):
            stripped = line.strip()
            if not stripped:
                continue

            header_match = re.match(r'^(#{1,4})\s+(.+)$', stripped)
            if header_match:
                level = min(len(header_match.group(1)), 2)
                content = header_match.group(2).replace('**', '').replace('_', '')
                doc.add_heading(content, level=level)
                continue

            if stripped.startswith(('* ', '- ', '• ')):
                item = re.sub(r'^[\*\-•]\s+', '', stripped)
                p = doc.add_paragraph(style='List Bullet')
                add_runs_with_bold(p, item)
                continue

            is_section = (
                stripped == stripped.upper()
                and len(stripped) > 3
                and '|' not in stripped
                and '@' not in stripped
                and not re.search(r'\d', stripped)
            )
            if is_section:
                doc.add_heading(stripped.replace('**', ''), level=2)
                continue

            clean_for_check = stripped.replace('**', '')
            if '|' in clean_for_check and clean_for_check.upper() == clean_for_check:
                p = doc.add_paragraph()
                add_runs_with_bold(p, stripped)
                for run in p.runs:
                    run.bold = True
                    run.font.size = Pt(11)
                continue

            if (
                re.match(r'^_.*_$', stripped)
                or re.match(r'^\d{2}/\d{2}/\d{4}', stripped)
                or re.search(r'(–|-)\s*(Current|Present)', stripped, re.I)
                or re.match(r'^(January|February|March|April|May|June|July|August|September|October|November|December)\s+\d{4}', stripped, re.I)
            ):
                p = doc.add_paragraph()
                clean_date = stripped.strip('_')
                run = p.add_run(clean_date.replace('**', ''))
                run.italic = True
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)
                continue

            p = doc.add_paragraph()
            add_runs_with_bold(p, stripped)

        file_stream = BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)

        print("✅ [DIAGNOZA EXPORT DOCX] Document generat cu succes.", flush=True)

        return send_file(
            file_stream,
            as_attachment=True,
            download_name="CV_Optimizat.docx",
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
    except Exception as e:
        print(f"❌ EROARE DOCX: {type(e).__name__} - {str(e)}", flush=True)
        traceback.print_exc()
        return api_response(error=f"Eroare generare DOCX: {str(e)}", code=500)


@app.route("/export-pdf", methods=["POST", "OPTIONS"], endpoint="export_pdf_root")
@app.route("/api/export-pdf", methods=["POST", "OPTIONS"], endpoint="export_pdf_api")
def export_pdf():
    if request.method == "OPTIONS":
        return "", 200

    try:
        data = request.get_json(force=True, silent=True) or {}
        text_content = data.get("text") or MEMORY.get("cv_text") or ""

        print(f"--- [DIAGNOZA EXPORT PDF] --- Lungime text: {len(text_content)} caractere", flush=True)

        if not text_content:
            return api_response(error="Text lipsa pentru export PDF.", code=400)

        buffer = BytesIO()
        doc = SimpleDocTemplate(
            buffer,
            pagesize=A4,
            rightMargin=16*mm,
            leftMargin=16*mm,
            topMargin=14*mm,
            bottomMargin=14*mm
        )

        styles = getSampleStyleSheet()

        style_name = ParagraphStyle(
            'CVName',
            parent=styles['Heading1'],
            fontName='Helvetica-Bold',
            fontSize=14,
            leading=18,
            spaceAfter=4,
            textColor='#0f172a',
        )
        style_heading = ParagraphStyle(
            'CVHeading',
            parent=styles['Heading2'],
            fontName='Helvetica-Bold',
            fontSize=11,
            leading=14,
            spaceBefore=12,
            spaceAfter=5,
            textColor='#0f172a',
            borderPadding=2,
        )
        style_job = ParagraphStyle(
            'CVJob',
            parent=styles['Normal'],
            fontName='Helvetica-Bold',
            fontSize=10.5,
            leading=13,
            spaceBefore=8,
            spaceAfter=1,
            textColor='#0f172a',
        )
        style_date = ParagraphStyle(
            'CVDate',
            parent=styles['Normal'],
            fontName='Helvetica-Oblique',
            fontSize=9,
            leading=11,
            spaceAfter=4,
            textColor='#64748b',
        )
        style_normal = ParagraphStyle(
            'CVNormal',
            parent=styles['Normal'],
            fontName='Helvetica',
            fontSize=10,
            leading=13,
            spaceAfter=3,
            textColor='#1e293b',
        )
        style_bullet = ParagraphStyle(
            'CVBullet',
            parent=styles['Normal'],
            fontName='Helvetica',
            fontSize=10,
            leading=13,
            leftIndent=12,
            spaceAfter=2,
            textColor='#334155',
        )

        def md_to_reportlab(text):
            """Converteste **bold** si _italic_ in tag-uri reportlab."""
            t = (
                text.replace('&', '&amp;')
                    .replace('<', '&lt;')
                    .replace('>', '&gt;')
            )
            t = re.sub(r'\*\*(.+?)\*\*', r'<b>\1</b>', t)
            t = re.sub(r'__(.+?)__', r'<b>\1</b>', t)
            t = re.sub(r'(?<!\*)\*(?!\*)(.+?)(?<!\*)\*(?!\*)', r'<i>\1</i>', t)
            t = re.sub(r'(?<!_)_(?!_)(.+?)(?<!_)_(?!_)', r'<i>\1</i>', t)
            return t

        story = []
        first_line = True

        for line in text_content.split("\n"):
            stripped = line.strip()
            if not stripped:
                story.append(Spacer(1, 3))
                continue

            header_match = re.match(r'^(#{1,4})\s+(.+)$', stripped)
            if header_match:
                level = len(header_match.group(1))
                content = md_to_reportlab(header_match.group(2))
                if level <= 2 and first_line:
                    story.append(Paragraph(content, style_name))
                    first_line = False
                else:
                    story.append(Paragraph(content, style_heading))
                continue

            if stripped.startswith(('* ', '- ', '• ')):
                item = re.sub(r'^[\*\-•]\s+', '', stripped)
                story.append(Paragraph('• ' + md_to_reportlab(item), style_bullet))
                continue

            clean = stripped.replace('**', '')
            if (clean == clean.upper() and len(clean) > 3
                    and '|' not in clean and '@' not in clean
                    and not re.search(r'\d', clean)):
                story.append(Paragraph(md_to_reportlab(clean), style_heading))
                continue

            if '|' in clean and clean.upper() == clean:
                story.append(Paragraph(md_to_reportlab(stripped), style_job))
                continue

            if (re.match(r'^_.*_$', stripped)
                    or re.match(r'^\d{2}/\d{2}/\d{4}', stripped)
                    or re.search(r'(–|-)\s*(Current|Present)', stripped, re.I)
                    or re.match(
                        r'^(January|February|March|April|May|June|July|August|September|October|November|December)\s+\d{4}',
                        stripped, re.I)):
                clean_date = stripped.strip('_')
                story.append(Paragraph(md_to_reportlab(clean_date), style_date))
                continue

            story.append(Paragraph(md_to_reportlab(stripped), style_normal))
            first_line = False

        doc.build(story)
        buffer.seek(0)

        print("✅ [DIAGNOZA EXPORT PDF] PDF generat cu succes.", flush=True)

        return send_file(
            buffer,
            as_attachment=True,
            download_name="CV_Optimizat.pdf",
            mimetype="application/pdf",
        )

    except Exception as e:
        print(f"❌ EROARE PDF: {type(e).__name__} - {str(e)}", flush=True)
        traceback.print_exc()
        return api_response(error=f"Eroare generare PDF: {str(e)}", code=500)


@app.route("/get-session", methods=["GET", "OPTIONS"], endpoint="get_session_root")
@app.route("/api/get-session", methods=["GET", "OPTIONS"], endpoint="get_session_api")
def get_session():
    if request.method == "OPTIONS":
        return api_response(code=200)
    return api_response(
        payload={
            "has_cv": bool(MEMORY.get("cv_text")),
            "has_job": bool(MEMORY.get("job_description")),
            "cv_length": len(MEMORY.get("cv_text", "")),
        }
    )


if __name__ == "__main__":
    app.run(host="0.0.0.0", port=int(os.environ.get("PORT", 5000)))
